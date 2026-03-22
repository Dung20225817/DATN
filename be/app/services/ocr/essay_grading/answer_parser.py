# essay_grading/answer_parser.py

import re
import json
from docx import Document
from typing import Dict, List, Tuple, Optional
import pdfplumber


def extract_rubric(doc: Document) -> Dict[int, dict]:
    """
    Extract RUBRIC section from DOCX
    
    Format:
    -------
    Câu 1 - Nghị luận văn học (5 điểm)
    Nội dung (2 điểm):
      - Nêu được một số nét đặc sắc [1 điểm]
      - Có những lí lẽ xác đáng [1 điểm]
    Hình thức (2 điểm):
      - Lập luận chặt chẽ [1 điểm]
      - Bố cục ba phần [1 điểm]
    
    Output:
    -------
    {
        1: {
            'question_text': 'Câu 1 - Nghị luận văn học',
            'total_points': 5,
            'criteria': {
                'Nội dung': {'points': 2, 'details': ['Nêu được...', 'Có những...']},
                'Hình thức': {'points': 2, 'details': ['Lập luận...', 'Bố cục...']},
                'Diễn đạt': {'points': 1, 'details': [...]}
            }
        },
        ...
    }
    """
    
    rubric = {}
    current_question = None
    current_criterion = None
    current_criterion_points = 0
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Extract question number and points: "Câu 1 - ... (5 điểm)"
        match_question = re.search(r'Câu\s+(\d+).*\((\d+)\s*điểm\)', text)
        if match_question:
            q_num = int(match_question.group(1))
            q_points = int(match_question.group(2))
            
            current_question = q_num
            rubric[q_num] = {
                'question_text': text,
                'total_points': q_points,
                'criteria': {}
            }
            current_criterion = None
            continue
        
        if current_question is None:
            continue
        
        # Extract criterion and points: "Nội dung (2 điểm):"
        match_criterion = re.search(r'([^(]+)\s*\((\d+)\s*điểm\)\s*:?', text)
        if match_criterion and ':' in text:
            criterion_name = match_criterion.group(1).strip()
            criterion_points = int(match_criterion.group(2))
            
            current_criterion = criterion_name
            rubric[current_question]['criteria'][criterion_name] = {
                'points': criterion_points,
                'details': []
            }
            continue
        
        # Extract detail points (bullet points)
        if current_criterion and (text.startswith('-') or text.startswith('•') or text.startswith('+')):
            detail = re.sub(r'^[-•+]\s*', '', text).strip()
            rubric[current_question]['criteria'][current_criterion]['details'].append(detail)
    
    return rubric


def extract_sample_answers(doc: Document) -> Dict[int, str]:
    """
    Extract SAMPLE_ANSWER section from DOCX
    
    Format:
    -------
    SAMPLE_ANSWER - MẪU ĐÁP ÁN
    Câu 1: [Title]
    [Full answer paragraph 1]
    [Full answer paragraph 2]
    ...
    
    Câu 2: [Title]
    [Full answer paragraph 1]
    ...
    
    Output:
    -------
    {
        1: "Full text of answer for question 1",
        2: "Full text of answer for question 2",
        ...
    }
    """
    
    samples = {}
    current_question = None
    answer_buffer = []
    in_sample_section = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            if answer_buffer and current_question is not None:
                # Save current answer before empty line
                samples[current_question] = ' '.join(answer_buffer).strip()
                answer_buffer = []
            continue
        
        # Detect SAMPLE_ANSWER section header
        if 'SAMPLE_ANSWER' in text.upper() or 'MAU DAP AN' in text.upper():
            in_sample_section = True
            continue
        
        if not in_sample_section:
            continue
        
        # Detect question marker: "Câu 1:", "Câu 2:", etc.
        match_question = re.match(r'C(ầ|au)\s+(\d+)\s*:', text)
        if match_question:
            # Save previous answer if exists
            if current_question is not None and answer_buffer:
                samples[current_question] = ' '.join(answer_buffer).strip()
            
            current_question = int(match_question.group(2))
            answer_buffer = []
            
            # Extract text after "Câu X: "
            remaining_text = re.sub(r'C(ầ|au)\s+\d+\s*:\s*', '', text)
            if remaining_text:
                answer_buffer.append(remaining_text)
            continue
        
        # Collect answer text
        if current_question is not None:
            answer_buffer.append(text)
    
    # Don't forget last answer
    if current_question is not None and answer_buffer:
        samples[current_question] = ' '.join(answer_buffer).strip()
    
    return samples


def extract_keywords(doc: Document) -> Dict[int, List[str]]:
    """
    Extract KEYWORDS section from DOCX
    
    Format:
    -------
    KEYWORDS - TỪ KHÓA PHẢI CÓ
    Câu 1:
      - Keyword 1
      - Keyword 2
      ...
    
    Câu 2:
      - Keyword 1
      ...
    
    Output:
    -------
    {
        1: ['Keyword 1', 'Keyword 2', ...],
        2: ['Keyword 1', ...],
        ...
    }
    """
    
    keywords = {}
    current_question = None
    in_keyword_section = False
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            continue
        
        # Detect KEYWORDS section header
        if 'KEYWORDS' in text.upper() or 'TU KHOA' in text.upper():
            in_keyword_section = True
            continue
        
        if not in_keyword_section:
            continue
        
        # Detect question marker: "Câu 1:", "Câu 2:", etc.
        match_question = re.match(r'C(ầ|au)\s+(\d+)\s*:', text)
        if match_question:
            current_question = int(match_question.group(2))
            keywords[current_question] = []
            continue
        
        # Extract keyword (bullet point)
        if current_question is not None and (text.startswith('-') or text.startswith('•')):
            keyword = re.sub(r'^[-•]\s*', '', text).strip()
            if keyword:
                keywords[current_question].append(keyword)
    
    return keywords


# All valid criteria sections supported
VALID_CRITERIA = ["Nội dung", "Hình thức", "Bố cục"]


def extract_criteria_format(doc: Document, criteria_list: List[str]) -> Dict[int, dict]:
    """
    Parse structured answer key format:

        Câu 1:
        Nội dung:
        + Ý 1 của nội dung
        + Ý 2 của nội dung
        Hình thức:
        Lập luận chặt chẽ; kết hợp thao tác nghị luận; diễn đạt mạch lạc
        Bố cục:
        Mở bài / Thân bài / Kết bài

    Rules:
    - Bullet lines (+, -, •): each line = one ý
    - Non-bullet lines: split by ";" to get individual ý
    - Points distributed equally across all ý in the question

    Returns:
    --------
    {
        1: {
            'question_text': 'Câu 1',
            'total_points': 10,
            'criteria': {
                'Nội dung': {'points': 4.0, 'details': ['ý 1', 'ý 2']},
                'Hình thức': {'points': 6.0, 'details': ['ý 1', 'ý 2', 'ý 3']},
            }
        }
    }
    """
    result: Dict[int, dict] = {}
    current_question: Optional[int] = None
    current_criterion: Optional[str] = None
    criterion_buffer: List[tuple] = []  # list of (text, is_bullet)

    criteria_map = {c.lower(): c for c in criteria_list}

    def flush():
        """Flush criterion_buffer into result."""
        if current_question is None or current_criterion is None or not criterion_buffer:
            return
        all_points = []
        for text, is_bullet in criterion_buffer:
            if is_bullet:
                clean = re.sub(r'^[\s+\-•]+', '', text).strip()
                if clean and len(clean) > 3:
                    all_points.append(clean)
            else:
                for part in re.split(r'\s*;\s*', text):
                    part = part.strip().rstrip('.')
                    if part and len(part) > 3:
                        all_points.append(part)
        result[current_question]['criteria'][current_criterion]['details'] = all_points

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # --- Detect question header: "Câu X" or "Câu X:" alone on the line ---
        q_match = re.match(r'^[Cc]âu\s*(\d+)\s*:?\s*$', text)
        if q_match:
            flush()
            current_question = int(q_match.group(1))
            current_criterion = None
            criterion_buffer = []
            result[current_question] = {
                'question_text': f'Câu {current_question}',
                'total_points': 10,
                'criteria': {}
            }
            continue

        if current_question is None:
            continue

        # --- Detect criterion heading: "Nội dung:", "Hình thức: text", etc. ---
        matched_c = None
        remainder = ''
        for c_lower, c_orig in criteria_map.items():
            if re.match(rf'^{re.escape(c_orig)}\s*:?\s*$', text, re.IGNORECASE):
                matched_c = c_orig
                remainder = ''
                break
            m = re.match(rf'^{re.escape(c_orig)}\s*:\s*(.+)$', text, re.IGNORECASE)
            if m:
                matched_c = c_orig
                remainder = m.group(1).strip()
                break

        if matched_c:
            flush()
            current_criterion = matched_c
            criterion_buffer = []
            result[current_question]['criteria'][matched_c] = {'points': 0, 'details': []}
            if remainder:
                criterion_buffer.append((remainder, remainder[0] in '+•-'))
            continue

        # --- Accumulate text for current criterion ---
        if current_criterion is not None:
            is_bullet = bool(text) and text[0] in '+•-'
            criterion_buffer.append((text, is_bullet))

    flush()  # Flush the last criterion

    # Distribute points equally per ý across all criteria in a question
    for q_data in result.values():
        total_ý = sum(len(c['details']) for c in q_data['criteria'].values())
        if total_ý > 0:
            point_per_ý = q_data['total_points'] / total_ý
            for c_data in q_data['criteria'].values():
                c_data['points'] = round(len(c_data['details']) * point_per_ý, 2)

    return result


def extract_flexible(doc: Document) -> Dict[str, str]:
    """
    Fallback parser: Extract questions and answers from any Vietnamese docx format.
    Detects "Câu X" or "Câu X:" patterns and collects text between them.
    
    Returns:
    --------
    { "Câu 1": "full answer text", "Câu 2": "...", ... }
    """
    samples = {}
    current_question = None
    buffer = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Match "Câu 1", "Câu 1:", "Câu 1 .", "Câu1:" etc.
        match = re.match(r'[Cc]âu\s*(\d+)\s*[:\.\-]?\s*(.*)', text)
        if match:
            if current_question and buffer:
                samples[current_question] = ' '.join(buffer).strip()
            current_question = f"Câu {match.group(1)}"
            remaining = match.group(2).strip()
            buffer = [remaining] if remaining else []
        elif current_question:
            buffer.append(text)

    if current_question and buffer:
        samples[current_question] = ' '.join(buffer).strip()

    return samples


def parse_answer_key_docx(docx_path: str, criteria_list: List[str] = None) -> Dict:
    """
    Main function: Parse DOCX answer key file.
    Tries strict format first (RUBRIC/SAMPLE_ANSWER sections),
    falls back to flexible extraction if strict format not found.
    
    Returns:
    --------
    {
        'status': 'success' | 'warning' | 'error',
        'message': str,
        'data': {
            'rubric': {...},
            'sample_answers': {...},
            'keywords': {...},
            'total_points': float,
            'num_questions': int,
            'questions': [1, 2, 3, ...]
        }
    }
    """
    
    try:
        doc = Document(docx_path)

        # --- CRITERIA FORMAT: user-selected criteria sections ---
        if criteria_list:
            criteria_rubric = extract_criteria_format(doc, criteria_list)

            if not criteria_rubric:
                return {
                    'status': 'error',
                    'message': (
                        f'Không tìm thấy câu nào với các tiêu chí [{", ".join(criteria_list)}]. '
                        f'File cần có các đoạn "Câu 1:", "Câu 2:" và các mục '
                        f'{" / ".join(criteria_list)} bên dưới mỗi câu.'
                    ),
                    'data': None
                }

            # Build flat sample_answers from all ý for backward-compat grading
            samples: Dict[int, str] = {}
            for q_num, q_data in criteria_rubric.items():
                all_details: List[str] = []
                for c_data in q_data['criteria'].values():
                    all_details.extend(c_data['details'])
                samples[q_num] = '. '.join(all_details)

            num_questions = len(criteria_rubric)
            total_points = sum(q['total_points'] for q in criteria_rubric.values())
            question_ids = sorted(criteria_rubric.keys())

            criteria_str = ', '.join(criteria_list)
            return {
                'status': 'success',
                'message': (
                    f'Phân tích theo tiêu chí [{criteria_str}]: '
                    f'{num_questions} câu, {total_points} điểm'
                ),
                'data': {
                    'rubric': criteria_rubric,
                    'sample_answers': samples,
                    'keywords': {},
                    'total_points': total_points,
                    'num_questions': num_questions,
                    'questions': question_ids
                }
            }

        # --- CLASSIC FORMAT: RUBRIC / SAMPLE_ANSWER / KEYWORDS sections ---
        # Extract all three sections
        rubric = extract_rubric(doc)
        samples = extract_sample_answers(doc)
        keywords = extract_keywords(doc)
        
        if not rubric:
            # Fallback: Try flexible extraction
            flexible_samples = extract_flexible(doc)
            
            if not flexible_samples:
                return {
                    'status': 'error',
                    'message': (
                        'Không tìm thấy câu hỏi nào trong file. '
                        'File cần có ít nhất các đoạn bắt đầu bằng "Câu 1", "Câu 2", v.v.'
                    ),
                    'data': None
                }
            
            # Build minimal rubric from flexible extraction
            for q_label in flexible_samples:
                q_num = int(re.search(r'\d+', q_label).group())
                rubric[q_num] = {
                    'question_text': q_label,
                    'total_points': 10,
                    'criteria': {'Nội dung': {'points': 10, 'details': []}}
                }
                samples[q_num] = flexible_samples[q_label]
            
            num_questions = len(rubric)
            total_points = sum(q['total_points'] for q in rubric.values())
            question_ids = sorted(rubric.keys())
            
            return {
                'status': 'warning',
                'message': (
                    f'Phân tích linh hoạt: tìm thấy {num_questions} câu. '
                    'Điểm mặc định 10đ/câu. '
                    'Để chấm điểm chi tiết, hãy dùng đúng format có phần RUBRIC.'
                ),
                'data': {
                    'rubric': rubric,
                    'sample_answers': samples,
                    'keywords': keywords,
                    'total_points': total_points,
                    'num_questions': num_questions,
                    'questions': question_ids
                }
            }
        
        # Calculate total points
        total_points = sum(q_data['total_points'] for q_data in rubric.values())
        num_questions = len(rubric)
        question_ids = sorted(rubric.keys())
        
        # Verify consistency
        if len(samples) != len(rubric):
            return {
                'status': 'warning',
                'message': f'Số câu trong RUBRIC ({len(rubric)}) và SAMPLE_ANSWER ({len(samples)}) không khớp',
                'data': {
                    'rubric': rubric,
                    'sample_answers': samples,
                    'keywords': keywords,
                    'total_points': total_points,
                    'num_questions': num_questions,
                    'questions': question_ids
                }
            }
        
        return {
            'status': 'success',
            'message': f'Được phân tích thành công: {num_questions} câu, {total_points} điểm',
            'data': {
                'rubric': rubric,
                'sample_answers': samples,
                'keywords': keywords,
                'total_points': total_points,
                'num_questions': num_questions,
                'questions': question_ids
            }
        }
    
    except Exception as e:
        return {
            'status': 'error',
            'message': f'Lỗi khi đọc file DOCX: {str(e)}',
            'data': None
        }


def parse_answer_key_pdf(pdf_path: str) -> Dict:
    """
    Parse PDF answer key file (Alternative to DOCX)
    
    Returns same format as parse_answer_key_docx()
    """
    
    try:
        text_content = ""
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text_content += page.extract_text() + "\n"
        
        # Parse text to extract RUBRIC, SAMPLE_ANSWER, KEYWORDS
        # This is more complex as we need to parse plain text
        # For now, return error suggesting to use DOCX format
        
        return {
            'status': 'warning',
            'message': 'PDF parsing chưa được hỗ trợ đầy đủ. Vui lòng sử dụng định dạng DOCX',
            'data': None
        }
    
    except Exception as e:
        return {
            'status': 'error',
            'message': f'Lỗi khi đọc file PDF: {str(e)}',
            'data': None
        }


def parse_answer_key(file_path: str, criteria_list: List[str] = None) -> Dict:
    """
    Wrapper function: Auto-detect file format and parse.

    Args:
        file_path: Path to .docx / .doc / .pdf file
        criteria_list: Optional list from VALID_CRITERIA e.g. ["Nội dung", "Hình thức"].
                       When provided, uses extract_criteria_format() instead of the
                       classic RUBRIC/SAMPLE_ANSWER/KEYWORDS parser.
    """

    if file_path.endswith('.docx') or file_path.endswith('.doc'):
        return parse_answer_key_docx(file_path, criteria_list)
    elif file_path.endswith('.pdf'):
        return parse_answer_key_pdf(file_path)
    else:
        return {
            'status': 'error',
            'message': 'Định dạng file không được hỗ trợ. Vui lòng sử dụng DOCX hoặc PDF',
            'data': None
        }
